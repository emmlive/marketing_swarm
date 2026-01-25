# ===========================
# SwarmDigiz — app.py (Robust)
# ===========================
import os
import json
import time
import sqlite3
from io import BytesIO
from datetime import datetime
from typing import Dict, Any, List, Tuple, Optional

import streamlit as st
import pandas as pd
import streamlit_authenticator as stauth
from docx import Document
from fpdf import FPDF

from main import run_marketing_swarm

# ============================================================
# APP
# ============================================================
APP_NAME = "SwarmDigiz"
DB_PATH = "breatheeasy.db"
APP_LOGO_PATH = "Logo1.jpeg"

PLAN_SEATS = {"Lite": 1, "Basic": 1, "Pro": 5, "Enterprise": 20, "Unlimited": 9999}
PLAN_AGENT_LIMITS = {"Lite": 3, "Basic": 3, "Pro": 5, "Enterprise": 8, "Unlimited": 12}

st.set_page_config(page_title=APP_NAME, layout="wide", initial_sidebar_state="expanded")

# ============================================================
# SESSION INIT
# ============================================================
def ss_init(key: str, default):
    if key not in st.session_state:
        st.session_state[key] = default

ss_init("theme_mode", "Night")
ss_init("sidebar_compact", False)
ss_init("notify_on_done", True)

ss_init("biz_name", "")
ss_init("directives", "")
ss_init("website_url", "")

# Swarm runner state
ss_init("swarm_running", False)
ss_init("swarm_paused", False)
ss_init("swarm_stop", False)
ss_init("swarm_autorun", True)
ss_init("swarm_autodelay", 3)     # 1/3/5
ss_init("swarm_next_ts", 0.0)
ss_init("swarm_queue", [])
ss_init("swarm_idx", 0)
ss_init("swarm_payload", {})
ss_init("last_active_swarm", [])

# outputs
ss_init("report", {})
ss_init("gen", False)

# Navigation helper (works even if tabs are not “switchable” programmatically)
ss_init("nav_choice", "🏠 Dashboard")

# ============================================================
# AGENTS (Seats)
# ============================================================
AGENT_UI: List[Tuple[str, str]] = [
    ("🕵️ Analyst", "analyst"),
    ("🧭 Marketing Adviser", "marketing_adviser"),
    ("🔎 Market Researcher", "market_researcher"),
    ("🛒 E-Commerce Marketer", "ecommerce_marketer"),
    ("📺 Ads", "ads"),
    ("🎨 Creative", "creative"),
    ("🧩 Guest Posting", "guest_posting"),
    ("👔 Strategist", "strategist"),
    ("📱 Social", "social"),
    ("📍 GEO", "geo"),
    ("📍 GBP Growth", "gbp_growth"),
    ("🌐 Website Audit", "audit"),
    ("✍ SEO", "seo"),
]

AGENT_SPECS: Dict[str, str] = {
    "analyst": "Competitor gaps, pricing, positioning, quick wins.",
    "marketing_adviser": "Messaging, channel priorities, measurement, next steps.",
    "market_researcher": "Segments, competitors, demand themes (no fake stats).",
    "ecommerce_marketer": "Offers, landing pages, flows, remarketing.",
    "ads": "Google/Meta ad copy tables + angles.",
    "creative": "Concepts + prompt packs + ad variants.",
    "guest_posting": "Outreach templates + PR/link plan.",
    "strategist": "30-day execution roadmap + KPIs.",
    "social": "30-day social plan with hooks & CTAs.",
    "geo": "Local visibility, citations, GBP checklist.",
    "gbp_growth": "Weekly GBP posts, review replies, keywords, ranking-drop triage.",
    "audit": "Conversion friction audit + fixes (requires website URL).",
    "seo": "Authority article + content clusters.",
}

DEPLOY_PROTOCOL = [
    "Configure mission: Brand + Location + Directives + Website URL (for Audit).",
    "Pick unlocked agents for this mission (locked agents are disabled).",
    "Launch Swarm. Agents run sequentially; you can Pause/Stop.",
    "Review each seat, refine, export Word/PDF, and save to Reports Vault.",
    "Manage execution in Team Intel (Projects + Kanban) and share with your team.",
]

SOCIAL_PUSH_PLATFORMS = [
    ("Google Business Profile", "https://business.google.com/"),
    ("Meta Business Suite (FB/IG)", "https://business.facebook.com/"),
    ("LinkedIn Company Page", "https://www.linkedin.com/company/"),
    ("X (Twitter)", "https://twitter.com/compose/tweet"),
]

# ============================================================
# THEME CSS (Night/Day + compact spacing + animation)
# ============================================================
def inject_theme_css():
    mode = st.session_state.get("theme_mode", "Night")
    compact = bool(st.session_state.get("sidebar_compact", False))

    if mode == "Night":
        bg = "#060B16"
        sidebar_bg = "#0B1220"
        panel = "rgba(15, 23, 42, 0.88)"
        text = "#F3F6FF"
        muted = "#AAB4C6"
        border = "rgba(148, 163, 184, 0.28)"
        input_bg = "rgba(17, 24, 39, 0.96)"
        input_text = "#F3F6FF"
        accent = "rgba(99,102,241,0.65)"
    else:
        bg = "#FFFFFF"
        sidebar_bg = "#F7FAFF"
        panel = "rgba(255,255,255,0.96)"
        text = "#0F172A"
        muted = "#475569"
        border = "rgba(15,23,42,0.14)"
        input_bg = "#FFFFFF"
        input_text = "#0F172A"
        accent = "rgba(99,102,241,0.30)"

    sbw = "260px" if compact else "360px"
    widget_gap = "0.35rem" if compact else "0.65rem"
    pad = "6px 10px" if compact else "12px 14px"

    st.markdown(f"""
    <style>
      .stApp {{
        background:
          radial-gradient(1200px 600px at 50% 0%, rgba(99,102,241,0.14), transparent 60%),
          linear-gradient(180deg, {bg} 0%, {bg} 100%);
      }}

      .stApp, .stApp p, .stApp li, .stApp h1, .stApp h2, .stApp h3, .stApp h4 {{
        color: {text} !important;
      }}
      .ms-muted {{ color: {muted} !important; }}

      .ms-card {{
        border: 1px solid {border};
        border-radius: 16px;
        background: {panel};
        padding: 14px 16px;
        box-shadow: 0 18px 55px rgba(0,0,0,0.28);
        margin-bottom: 10px;
      }}
      .ms-chip {{
        display:inline-block;
        padding:4px 10px;
        border-radius:999px;
        border:1px solid {accent};
        background: rgba(99,102,241,0.14);
        color: {text} !important;
        font-size:12px;
      }}

      /* inputs visible */
      .stTextInput input,
      .stTextArea textarea {{
        background: {input_bg} !important;
        color: {input_text} !important;
        border: 1px solid {border} !important;
        border-radius: 10px !important;
      }}
      .stTextInput input::placeholder,
      .stTextArea textarea::placeholder {{
        color: {muted} !important;
        opacity: .9 !important;
      }}
      div[data-baseweb="select"] > div {{
        background: {input_bg} !important;
        color: {input_text} !important;
        border: 1px solid {border} !important;
        border-radius: 10px !important;
      }}

      section[data-testid="stSidebar"] {{
        width: {sbw} !important;
        transition: width 240ms ease-in-out;
      }}
      section[data-testid="stSidebar"] > div {{
        background: {sidebar_bg};
        border-right: 1px solid {border};
        padding: {pad};
      }}

      section[data-testid="stSidebar"] .stMarkdown,
      section[data-testid="stSidebar"] .stTextInput,
      section[data-testid="stSidebar"] .stTextArea,
      section[data-testid="stSidebar"] .stSelectbox,
      section[data-testid="stSidebar"] .stRadio,
      section[data-testid="stSidebar"] .stToggle,
      section[data-testid="stSidebar"] .stButton,
      section[data-testid="stSidebar"] .stCheckbox {{
        margin-bottom: {widget_gap} !important;
      }}

      .stTabs [data-baseweb="tab-list"] {{ gap: 6px; }}

      #MainMenu {{visibility:hidden;}}
      footer {{visibility:hidden;}}
    </style>
    """, unsafe_allow_html=True)

inject_theme_css()

# ============================================================
# DB + HELPERS
# ============================================================
def db_conn() -> sqlite3.Connection:
    return sqlite3.connect(DB_PATH, check_same_thread=False)

def ensure_column(conn: sqlite3.Connection, table: str, col: str, col_def: str):
    cur = conn.cursor()
    cur.execute(f"PRAGMA table_info({table})")
    cols = {row[1] for row in cur.fetchall()}
    if col not in cols:
        cur.execute(f"ALTER TABLE {table} ADD COLUMN {col} {col_def}")

def _hash_password(pw: str) -> str:
    pw = pw or ""
    try:
        return stauth.Hasher([pw]).generate()[0]
    except Exception:
        return pw

@st.cache_resource
def init_db_once():
    conn = db_conn()
    cur = conn.cursor()

    cur.execute("""
        CREATE TABLE IF NOT EXISTS orgs (
            team_id TEXT PRIMARY KEY,
            org_name TEXT,
            plan TEXT DEFAULT 'Lite',
            seats_allowed INTEGER DEFAULT 1,
            allowed_agents_json TEXT DEFAULT '',
            status TEXT DEFAULT 'active',
            created_at TEXT DEFAULT (datetime('now'))
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS users (
            username TEXT PRIMARY KEY,
            email TEXT,
            name TEXT,
            password TEXT,
            role TEXT DEFAULT 'viewer',
            active INTEGER DEFAULT 1,
            plan TEXT DEFAULT 'Lite',
            credits INTEGER DEFAULT 10,
            verified INTEGER DEFAULT 1,
            team_id TEXT DEFAULT 'ORG_001',
            created_at TEXT DEFAULT (datetime('now')),
            last_login_at TEXT
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS audit_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            timestamp TEXT,
            team_id TEXT,
            actor TEXT,
            actor_role TEXT,
            action_type TEXT,
            object_type TEXT,
            object_id TEXT,
            details TEXT
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS geo_locations (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            state TEXT,
            city TEXT,
            team_id TEXT DEFAULT ''
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS leads (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            team_id TEXT,
            title TEXT,
            city TEXT,
            service TEXT,
            stage TEXT DEFAULT 'Discovery',
            created_at TEXT DEFAULT (datetime('now'))
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS projects (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            team_id TEXT,
            name TEXT,
            owner TEXT,
            status TEXT DEFAULT 'Active',
            notes TEXT,
            created_at TEXT DEFAULT (datetime('now'))
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS reports_vault (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            team_id TEXT,
            name TEXT,
            created_by TEXT,
            location TEXT,
            biz_name TEXT,
            selected_agents_json TEXT,
            report_json TEXT,
            full_report TEXT,
            created_at TEXT DEFAULT (datetime('now'))
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS feedback (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            team_id TEXT,
            username TEXT,
            rating INTEGER,
            message TEXT,
            page TEXT,
            created_at TEXT DEFAULT (datetime('now'))
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS upgrade_requests (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            team_id TEXT,
            username TEXT,
            current_plan TEXT,
            requested_plan TEXT,
            notes TEXT,
            status TEXT DEFAULT 'pending',
            created_at TEXT DEFAULT (datetime('now'))
        )
    """)

    ensure_column(conn, "orgs", "allowed_agents_json", "TEXT DEFAULT ''")
    ensure_column(conn, "orgs", "seats_allowed", "INTEGER DEFAULT 1")

    # seed geo if empty
    cur.execute("SELECT COUNT(*) FROM geo_locations")
    if int(cur.fetchone()[0] or 0) == 0:
        seed = {
            "Alabama": ["Birmingham", "Huntsville", "Mobile"],
            "Illinois": ["Chicago", "Naperville", "Plainfield"],
            "Texas": ["Austin", "Dallas", "Houston"],
            "California": ["Los Angeles", "San Francisco", "San Diego"],
            "Florida": ["Miami", "Orlando", "Tampa"],
        }
        for stt, cities in seed.items():
            for c in cities:
                cur.execute("INSERT INTO geo_locations (state, city, team_id) VALUES (?,?,?)", (stt, c, ""))

    # root org/user
    cur.execute("""
        INSERT OR IGNORE INTO orgs (team_id, org_name, plan, seats_allowed, status, allowed_agents_json)
        VALUES ('ROOT', 'SaaS Root', 'Unlimited', 9999, 'active', '')
    """)
    root_pw = _hash_password(os.getenv("ROOT_PASSWORD", "root123"))
    cur.execute("""
        INSERT OR REPLACE INTO users
        (username,email,name,password,role,active,plan,credits,verified,team_id)
        VALUES ('root','root@swarmdigiz.ai','Root Admin',?, 'root', 1,'Unlimited',9999,1,'ROOT')
    """, (root_pw,))

    # demo org
    cur.execute("SELECT COUNT(*) FROM orgs WHERE team_id!='ROOT'")
    if int(cur.fetchone()[0] or 0) == 0:
        allowed = json.dumps([k for _, k in AGENT_UI][:3])
        cur.execute("""
            INSERT OR IGNORE INTO orgs (team_id, org_name, plan, seats_allowed, status, allowed_agents_json)
            VALUES ('ORG_001','TechNovance Customer','Lite',1,'active',?)
        """, (allowed,))
        admin_pw = _hash_password("admin123")
        cur.execute("""
            INSERT OR REPLACE INTO users
            (username,email,name,password,role,active,plan,credits,verified,team_id)
            VALUES ('admin','admin@customer.ai','Org Admin',?, 'admin',1,'Lite',999,1,'ORG_001')
        """, (admin_pw,))

    conn.commit()
    conn.close()

init_db_once()

def log_audit(team_id: str, actor: str, role: str, action: str, obj_type="", obj_id="", details=""):
    try:
        conn = db_conn()
        conn.execute("""
            INSERT INTO audit_logs (timestamp,team_id,actor,actor_role,action_type,object_type,object_id,details)
            VALUES (?,?,?,?,?,?,?,?)
        """, (datetime.utcnow().isoformat(), team_id, actor, role, action, obj_type, obj_id, str(details)[:4000]))
        conn.commit(); conn.close()
    except Exception:
        pass

def get_user(username: str) -> Dict[str, Any]:
    conn = db_conn()
    df = pd.read_sql_query("SELECT * FROM users WHERE username=?", conn, params=(username,))
    conn.close()
    return df.iloc[0].to_dict() if not df.empty else {}

def get_org(team_id: str) -> Dict[str, Any]:
    conn = db_conn()
    df = pd.read_sql_query("SELECT * FROM orgs WHERE team_id=?", conn, params=(team_id,))
    conn.close()
    return df.iloc[0].to_dict() if not df.empty else {"team_id": team_id, "org_name": team_id, "plan": "Lite", "seats_allowed": 1}

def normalize_role(role: str) -> str:
    role = (role or "").strip().lower()
    return role if role in {"viewer","editor","admin","root"} else "viewer"

def active_user_count(team_id: str) -> int:
    conn = db_conn()
    df = pd.read_sql_query("SELECT COUNT(*) AS n FROM users WHERE team_id=? AND active=1 AND role!='root'", conn, params=(team_id,))
    conn.close()
    return int(df.iloc[0]["n"] or 0)

def seats_allowed_for_team(team_id: str) -> int:
    org = get_org(team_id)
    plan = str(org.get("plan", "Lite"))
    return int(org.get("seats_allowed") or PLAN_SEATS.get(plan, 1))

PERMISSIONS = {
    "viewer": {"read"},
    "editor": {"read"},
    "admin": {"read", "user_manage", "export", "project_manage"},
    "root": {"*"},
}
def can(role: str, perm: str) -> bool:
    perms = PERMISSIONS.get(normalize_role(role), {"read"})
    return "*" in perms or perm in perms or perm == "read"

def plan_agent_limit(plan: str) -> int:
    return int(PLAN_AGENT_LIMITS.get(plan, 3))

def default_allowed_agents_for_plan(plan: str) -> List[str]:
    keys = [k for _, k in AGENT_UI]
    return keys[:min(plan_agent_limit(plan), len(keys))]

def get_allowed_agents(team_id: str) -> List[str]:
    org = get_org(team_id)
    raw = (org.get("allowed_agents_json") or "").strip()
    try:
        lst = json.loads(raw) if raw else []
    except Exception:
        lst = []
    valid = {k for _, k in AGENT_UI}
    lst = [x for x in lst if x in valid]
    if lst:
        return lst
    auto = default_allowed_agents_for_plan(org.get("plan", "Lite"))
    conn = db_conn()
    conn.execute("UPDATE orgs SET allowed_agents_json=? WHERE team_id=?", (json.dumps(auto), team_id))
    conn.commit(); conn.close()
    return auto

# ============================================================
# EXPORT HELPERS
# ============================================================
def export_word(content: str, title: str) -> bytes:
    doc = Document()
    doc.add_heading(str(title), 0)
    for line in str(content).split("\n"):
        doc.add_paragraph(line)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

def export_pdf(content: str, title: str) -> bytes:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=14)
    pdf.add_page()
    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 10, str(title), ln=True)
    pdf.ln(4)
    pdf.set_font("Arial", size=10)
    safe_text = str(content).encode("latin-1", "ignore").decode("latin-1")
    pdf.multi_cell(0, 6, safe_text)
    return pdf.output(dest="S").encode("latin-1")

# ============================================================
# AUTH
# ============================================================
def get_db_creds():
    conn = db_conn()
    df = pd.read_sql_query("SELECT username,email,name,password FROM users WHERE active=1", conn)
    conn.close()
    return {
        "usernames": {
            r["username"]: {"email": r.get("email",""), "name": r.get("name", r["username"]), "password": r["password"]}
            for _, r in df.iterrows()
        }
    }

cookie_name = st.secrets.get("cookie", {}).get("name", "swarmdigiz_cookie")
cookie_key = st.secrets.get("cookie", {}).get("key", "swarmdigiz_cookie_key_change_me")
cookie_days = int(st.secrets.get("cookie", {}).get("expiry_days", 30))
authenticator = stauth.Authenticate(get_db_creds(), cookie_name, cookie_key, cookie_days)

def login_page():
    st.markdown(f"""
    <div class="ms-card">
      <h1 style="margin:0;">{APP_NAME}</h1>
      <div class="ms-muted">Root login: <b>root / root123</b> (or ROOT_PASSWORD env var).</div>
    </div>
    """, unsafe_allow_html=True)

    tabs = st.tabs(["Login", "Forgot Password", "Team Intel Login", "Pricing (placeholder)"])
    with tabs[0]:
        authenticator.login(location="main")
    with tabs[1]:
        authenticator.forgot_password(location="main")
    with tabs[2]:
        st.info("Team Intel is available after login. Your Org Admin provides username/password.")
        st.markdown("- Viewer: read-only\n- Admin: full org tools\n- Root: full SaaS tools")
    with tabs[3]:
        st.markdown("""
        <div class="ms-card">
          <b>Pricing (placeholder)</b> <span class="ms-chip">Move to Landing Page</span>
          <ul>
            <li><b>Lite</b> — 3 agents / 1 seat</li>
            <li><b>Pro</b> — 5 agents / 5 seats</li>
            <li><b>Enterprise</b> — 8+ agents / 20 seats</li>
          </ul>
        </div>
        """, unsafe_allow_html=True)

    st.stop()

if not st.session_state.get("authentication_status"):
    login_page()

# ============================================================
# CONTEXT
# ============================================================
me = get_user(st.session_state["username"])
my_team = me.get("team_id", "ORG_001")
my_role = normalize_role(me.get("role", "viewer"))
is_root = (my_team == "ROOT") or (my_role == "root")
org = get_org(my_team)
org_plan = str(org.get("plan", "Lite"))
unlocked_agents = [k for _, k in AGENT_UI] if is_root else get_allowed_agents(my_team)

# ============================================================
# Helpers for report + integrity + retry
# ============================================================
def is_placeholder(val: Any) -> bool:
    if val is None:
        return True
    s = str(val).strip().lower()
    return (not s) or s.startswith("agent not selected") or "no output returned" in s

def run_one(agent_key: str, payload: Dict[str, Any]) -> Dict[str, Any]:
    p = dict(payload)
    p["active_swarm"] = [agent_key]
    return run_marketing_swarm(p) or {}

def build_full_report(payload: Dict[str, Any], report: Dict[str, Any]) -> str:
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    head = f"# {payload.get('biz_name','')} Intelligence Report\n**Date:** {now} | **Location:** {payload.get('city','')} | **Plan:** {payload.get('package','')}\n---\n\n"
    parts = []
    for label, k in AGENT_UI:
        if k in report and not is_placeholder(report.get(k)):
            parts.append(f"## {label}\n{report.get(k)}")
    return head + ("\n\n".join(parts) if parts else "## Summary\nNo outputs generated.")

def report_integrity_df(report: Dict[str, Any], selected: List[str]) -> pd.DataFrame:
    rows = []
    for _, k in AGENT_UI:
        if k not in selected:
            continue
        v = report.get(k, "")
        rows.append({"agent": k, "status": "OK" if (v and not is_placeholder(v)) else "EMPTY", "chars": len(str(v or ""))})
    return pd.DataFrame(rows)

def retry_agent(agent_key: str):
    payload = dict(st.session_state.get("swarm_payload") or {})
    if not payload:
        st.error("No mission payload found. Launch a swarm first.")
        return
    with st.status(f"Retrying {agent_key}…", expanded=False):
        out = run_one(agent_key, payload)
    rep = dict(st.session_state.get("report") or {})
    if agent_key in out:
        rep[agent_key] = out.get(agent_key)
    rep["full_report"] = build_full_report(payload, rep)
    st.session_state["report"] = rep
    st.toast(f"✅ Retried {agent_key}", icon="✅")
    st.rerun()

# ============================================================
# SIDEBAR (Mission Control)
# ============================================================
with st.sidebar:
    badge = "🛡 ROOT" if is_root else ("👑 ORG ADMIN" if my_role == "admin" else "👁 VIEWER")
    st.markdown(f"## {APP_NAME} <span class='ms-chip'>{badge}</span>", unsafe_allow_html=True)
    st.caption(f"Team: `{my_team}` • Plan: **{org_plan}**")
    st.caption(f"Seats: {active_user_count(my_team)}/{seats_allowed_for_team(my_team)}")

    st.selectbox("🌗 Theme", ["Night", "Day"], index=0 if st.session_state["theme_mode"]=="Night" else 1, key="theme_mode")
    st.checkbox("🧩 Compact Sidebar", value=st.session_state["sidebar_compact"], key="sidebar_compact")
    inject_theme_css()

    st.divider()

    # Quick nav (works even while swarm is running)
    nav_items = ["🏠 Dashboard", "📖 Guide"] + [lbl for lbl, _ in AGENT_UI] + ["🤝 Team Intel"]
    if my_role in {"admin","root"}:
        nav_items.append("⚙ Org Admin")
    if is_root:
        nav_items.append("🛡 Root Admin")

    st.selectbox("🧭 Quick Nav", nav_items, key="nav_choice")

    st.divider()

    st.text_input("🏢 Brand Name", key="biz_name")
    st.text_input("🌐 Website URL (for Audit)", key="website_url")
    st.text_area("✍️ Strategic Directives", key="directives", height=90)

    # Dynamic Geo
    conn = db_conn()
    geo_df = pd.read_sql_query("SELECT state, city FROM geo_locations WHERE team_id IN ('', ?) ORDER BY state, city", conn, params=(my_team,))
    conn.close()
    states = sorted(list(geo_df["state"].unique()))
    state = st.selectbox("🎯 Target State", states)
    mode = st.radio("City", ["Pick from list", "Add custom"], horizontal=True, key="city_mode")
    if mode == "Pick from list":
        cities = sorted(list(geo_df[geo_df["state"] == state]["city"].unique()))
        city = st.selectbox("🏙️ Target City", cities)
    else:
        city = st.text_input("🏙️ Custom City", value="")
        if st.button("➕ Save City", use_container_width=True) and city.strip():
            conn = db_conn()
            conn.execute("INSERT INTO geo_locations (state, city, team_id) VALUES (?,?,?)", (state, city.strip(), my_team))
            conn.commit(); conn.close()
            st.toast("Saved city.", icon="✅")
            st.rerun()

    full_loc = f"{city}, {state}".strip(", ").strip()

    st.divider()

    st.checkbox("🔔 Notify when complete", key="notify_on_done")
    st.checkbox("⚡ Auto-run remaining agents", key="swarm_autorun")
    st.selectbox("⏱ Auto-run delay", [1,3,5], key="swarm_autodelay")

    with st.expander("🤖 Swarm Personnel", expanded=True):
        st.caption("Pick unlocked agents (locked agents grayed out).")
        for label, key in AGENT_UI:
            kk = f"tg_{key}"
            ss_init(kk, False)
            locked = (not is_root) and (key not in unlocked_agents)
            st.toggle(label, key=kk, disabled=locked)

    st.divider()

    # Dashboard jump while running
    if st.session_state["swarm_running"]:
        if st.button("🏠 Go to Dashboard", use_container_width=True):
            st.session_state["nav_choice"] = "🏠 Dashboard"
            st.rerun()

    if not st.session_state["swarm_running"]:
        if st.button("🚀 LAUNCH OMNI-SWARM", type="primary", use_container_width=True):
            selected = [k for _, k in AGENT_UI if st.session_state.get(f"tg_{k}", False)]
            if not st.session_state["biz_name"].strip():
                st.error("Enter Brand Name.")
            elif not selected:
                st.warning("Select at least one agent.")
            else:
                st.session_state["swarm_payload"] = {
                    "biz_name": st.session_state["biz_name"].strip(),
                    "city": full_loc or "USA",
                    "directives": st.session_state["directives"].strip(),
                    "url": st.session_state["website_url"].strip(),
                    "package": org_plan,
                }
                st.session_state["swarm_queue"] = selected[:]
                st.session_state["swarm_idx"] = 0
                st.session_state["swarm_next_ts"] = time.time()
                st.session_state["swarm_running"] = True
                st.session_state["swarm_paused"] = False
                st.session_state["swarm_stop"] = False
                st.session_state["report"] = {}
                st.session_state["gen"] = False
                st.session_state["last_active_swarm"] = selected[:]
                st.toast("Swarm started 🚀", icon="🚀")
                st.session_state["nav_choice"] = "🏠 Dashboard"
                st.rerun()
    else:
        c1,c2,c3 = st.columns(3)
        with c1:
            if st.button("⏸ Pause", use_container_width=True):
                st.session_state["swarm_paused"] = True
                st.rerun()
        with c2:
            if st.button("▶ Resume", use_container_width=True):
                st.session_state["swarm_paused"] = False
                st.session_state["swarm_next_ts"] = time.time()
                st.rerun()
        with c3:
            if st.button("🛑 Stop", use_container_width=True):
                st.session_state["swarm_stop"] = True
                st.session_state["swarm_running"] = False
                st.toast("Swarm stopped.", icon="🛑")
                st.session_state["nav_choice"] = "🏠 Dashboard"
                st.rerun()

    st.divider()

    # Feedback intake (always available)
    with st.expander("💬 Feedback", expanded=False):
        with st.form("feedback_form_sidebar"):
            rating = st.selectbox("Rating", [5,4,3,2,1], index=0)
            msg = st.text_area("Message", placeholder="What should we improve?")
            page = st.text_input("Page (optional)", value=st.session_state.get("nav_choice",""))
            submit = st.form_submit_button("Send", use_container_width=True)
        if submit and msg.strip():
            conn = db_conn()
            conn.execute("INSERT INTO feedback (team_id,username,rating,message,page) VALUES (?,?,?,?,?)",
                         (my_team, me.get("username",""), int(rating), msg.strip(), page.strip()))
            conn.commit(); conn.close()
            st.toast("Thanks! Feedback received ✅", icon="✅")

    authenticator.logout("🔒 Sign Out", "sidebar")

# ============================================================
# SWARM RUNNER (one-by-one + delay)
# ============================================================
if st.session_state["swarm_running"] and (not st.session_state["swarm_paused"]) and (not st.session_state["swarm_stop"]):
    if time.time() >= float(st.session_state["swarm_next_ts"]):
        q = st.session_state["swarm_queue"]
        idx = int(st.session_state["swarm_idx"])

        if idx >= len(q):
            st.session_state["swarm_running"] = False
            st.session_state["gen"] = True
            if st.session_state.get("notify_on_done", True):
                st.toast("✅ Swarm completed.", icon="✅")
            st.rerun()

        agent = q[idx]
        payload = dict(st.session_state["swarm_payload"])
        with st.status(f"Running {agent}…", expanded=False):
            try:
                out = run_one(agent, payload)
            except Exception as e:
                out = {agent: f"❌ Error: {e}"}

        rep = dict(st.session_state["report"] or {})
        if agent in out:
            rep[agent] = out.get(agent)
        rep["full_report"] = build_full_report(payload, rep)
        st.session_state["report"] = rep

        st.session_state["swarm_idx"] = idx + 1
        st.session_state["swarm_next_ts"] = time.time() + float(st.session_state["swarm_autodelay"]) if st.session_state["swarm_autorun"] else 10**12
        if not st.session_state["swarm_autorun"]:
            st.session_state["swarm_paused"] = True

        st.rerun()

# ============================================================
# PAGES (Quick Nav renders ONE page; avoids duplicate keys/forms)
# ============================================================
def render_dashboard():
    st.markdown(f"<div class='ms-card'><h2 style='margin:0;'>🏠 Dashboard</h2>"
                f"<div class='ms-muted'>Mission Control • Team `{my_team}` • Plan `{org_plan}`</div></div>",
                unsafe_allow_html=True)

    # Swarm status
    if st.session_state["swarm_running"]:
        q = st.session_state["swarm_queue"]
        idx = int(st.session_state["swarm_idx"])
        nxt = q[idx] if idx < len(q) else "—"
        st.markdown(f"<div class='ms-card'><b>🚀 Swarm Running</b> <span class='ms-chip'>Next: {nxt}</span>"
                    f"<div class='ms-muted'>Auto-run delay: {st.session_state['swarm_autodelay']}s • Auto-run: {st.session_state['swarm_autorun']}</div></div>",
                    unsafe_allow_html=True)
    else:
        st.markdown("<div class='ms-card'><b>Swarm Status:</b> Idle</div>", unsafe_allow_html=True)

    # Integrity check quick view
    selected = st.session_state.get("last_active_swarm", []) or []
    rep = st.session_state.get("report", {}) or {}
    if selected:
        st.subheader("Report Integrity Check")
        df = report_integrity_df(rep, selected)
        st.dataframe(df, use_container_width=True, hide_index=True)
        empty = [r["agent"] for r in df.to_dict("records") if r["status"] == "EMPTY"]
        if empty:
            st.warning(f"Empty outputs: {empty}")
            cols = st.columns(4)
            for i, a in enumerate(empty):
                with cols[i % 4]:
                    st.button(f"Retry {a}", key=f"dash_retry_{a}", on_click=retry_agent, args=(a,), use_container_width=True)
    else:
        st.info("Run a swarm to see integrity results.")

def render_guide():
    st.header(f"📖 {APP_NAME} Guide")
    st.markdown("### Protocol")
    for line in DEPLOY_PROTOCOL:
        st.markdown(f"- {line}")

def render_seat(label: str, key: str):
    st.subheader(f"{label} Seat")
    st.caption(AGENT_SPECS.get(key, ""))
    st.info(seat_how_to_use(key))

    rep = st.session_state.get("report", {}) or {}
    if key not in rep or is_placeholder(rep.get(key)):
        st.warning("No report yet for this seat. Select agent + run Swarm.")
        if key in (st.session_state.get("last_active_swarm") or []):
            st.button("🔁 Retry this agent", key=f"seat_retry_{key}", on_click=retry_agent, args=(key,))
        return

    edited = st.text_area("Refine Intel", value=str(rep.get(key)), height=380, key=f"ed_{key}")
    c1, c2, c3 = st.columns(3)
    with c1:
        st.download_button("📄 Word", export_word(edited, label), file_name=f"{key}.docx", key=f"w_{key}", use_container_width=True)
    with c2:
        st.download_button("📕 PDF", export_pdf(edited, label), file_name=f"{key}.pdf", key=f"p_{key}", use_container_width=True)
    with c3:
        st.button("🔁 Retry", key=f"retry_{key}", on_click=retry_agent, args=(key,), use_container_width=True)

def render_team_intel(prefix: str):
    """prefix is required to avoid duplicate form keys across views."""
    st.header("🤝 Team Intel")
    is_admin_like = (normalize_role(my_role) in {"admin","root"})
    tabs = st.tabs(["📌 Projects", "📋 Kanban Leads", "🧾 Reports Vault", "👥 Users & RBAC", "🔐 Security Logs"])

    def viewer_notice():
        st.info("Viewer: read-only. Ask Org Admin for edit access.")

    with tabs[0]:
        conn = db_conn()
        df = pd.read_sql_query("SELECT id,name,owner,status,created_at FROM projects WHERE team_id=? ORDER BY id DESC", conn, params=(my_team,))
        conn.close()
        st.dataframe(df, use_container_width=True, hide_index=True)

        if is_admin_like:
            with st.form(f"{prefix}_proj_new"):
                st.subheader("Create Project")
                name = st.text_input("Project name", key=f"{prefix}_proj_name")
                owner = st.text_input("Owner", value=me.get("username",""), key=f"{prefix}_proj_owner")
                status = st.selectbox("Status", ["Active","Paused","Done"], index=0, key=f"{prefix}_proj_status")
                notes = st.text_area("Notes", key=f"{prefix}_proj_notes")
                submit = st.form_submit_button("Create", use_container_width=True)
            if submit and name.strip():
                conn = db_conn()
                conn.execute("INSERT INTO projects (team_id,name,owner,status,notes) VALUES (?,?,?,?,?)",
                             (my_team,name.strip(),owner.strip(),status,notes.strip()))
                conn.commit(); conn.close()
                log_audit(my_team, me.get("username",""), my_role, "project.create", "project", "", name)
                st.toast("Project created ✅", icon="✅")
                st.rerun()
        else:
            viewer_notice()

    with tabs[1]:
        # drag-like kanban + bulk editor
        kanban_board(my_team, editable=is_admin_like)
        if not is_admin_like:
            viewer_notice()

    with tabs[2]:
        conn = db_conn()
        vdf = pd.read_sql_query("SELECT id,name,biz_name,location,created_by,created_at FROM reports_vault WHERE team_id=? ORDER BY id DESC", conn, params=(my_team,))
        conn.close()
        st.dataframe(vdf, use_container_width=True, hide_index=True)

        rep = st.session_state.get("report", {}) or {}
        if is_admin_like and rep:
            with st.form(f"{prefix}_vault_save"):
                name = st.text_input("Report name", value=f"{st.session_state.get('biz_name','Report')} • {datetime.now().strftime('%Y-%m-%d %H:%M')}", key=f"{prefix}_vault_name")
                submit = st.form_submit_button("Save Current Report", use_container_width=True)
            if submit:
                payload = st.session_state.get("swarm_payload", {}) or {}
                conn = db_conn()
                conn.execute("""
                    INSERT INTO reports_vault (team_id,name,created_by,location,biz_name,selected_agents_json,report_json,full_report)
                    VALUES (?,?,?,?,?,?,?,?)
                """, (my_team,name,me.get("username",""),payload.get("city",""),payload.get("biz_name",""),
                      json.dumps(st.session_state.get("last_active_swarm",[])), json.dumps(rep), rep.get("full_report","")))
                conn.commit(); conn.close()
                log_audit(my_team, me.get("username",""), my_role, "vault.save", "report", "", name)
                st.toast("Saved to Vault ✅", icon="✅")
                st.rerun()
        elif not is_admin_like:
            viewer_notice()

    with tabs[3]:
        conn = db_conn()
        udf = pd.read_sql_query("SELECT username,name,email,role,credits,active,last_login_at,created_at FROM users WHERE team_id=? AND role!='root' ORDER BY created_at DESC", conn, params=(my_team,))
        conn.close()
        st.dataframe(udf, use_container_width=True, hide_index=True)
        st.caption(f"Seats: {active_user_count(my_team)}/{seats_allowed_for_team(my_team)}")

        if is_admin_like:
            with st.form(f"{prefix}_add_user"):
                st.subheader("Add User")
                u = st.text_input("Username", key=f"{prefix}_u")
                n = st.text_input("Name", key=f"{prefix}_n")
                e = st.text_input("Email", key=f"{prefix}_e")
                r = st.selectbox("Role", ["viewer","editor","admin"], index=0, key=f"{prefix}_r")
                pw = st.text_input("Temp Password", type="password", key=f"{prefix}_pw")
                submit = st.form_submit_button("Create", use_container_width=True)
            if submit and u.strip() and pw.strip():
                conn = db_conn()
                conn.execute("INSERT INTO users (username,email,name,password,role,active,plan,credits,verified,team_id) VALUES (?,?,?,?,?,1,?,?,1,?)",
                             (u.strip(),e.strip(),n.strip(),_hash_password(pw.strip()),r,org_plan,10,my_team))
                conn.commit(); conn.close()
                log_audit(my_team, me.get("username",""), my_role, "user.create", "user", u, f"role={r}")
                st.toast("User created ✅", icon="✅")
                st.rerun()
        else:
            viewer_notice()

    with tabs[4]:
        conn = db_conn()
        logs = pd.read_sql_query("SELECT timestamp,actor,actor_role,action_type,object_type,object_id,details FROM audit_logs WHERE team_id=? ORDER BY id DESC LIMIT 250", conn, params=(my_team,))
        conn.close()
        st.dataframe(logs, use_container_width=True, hide_index=True)

def render_org_admin():
    st.header("⚙ Org Admin")
    st.caption("Org settings, upgrade requests, and team management.")
    st.markdown(f"<div class='ms-card'><b>Org:</b> {org.get('org_name','')} • <b>Team:</b> {my_team} • <b>Plan:</b> {org_plan}</div>", unsafe_allow_html=True)

    # Upgrade request
    st.subheader("Package Upgrade Request")
    with st.form("upgrade_request_form"):
        requested_plan = st.selectbox("Requested plan", ["Pro", "Enterprise", "Unlimited"], index=0)
        notes = st.text_area("Notes (optional)")
        submit = st.form_submit_button("Submit upgrade request", use_container_width=True)
    if submit:
        conn = db_conn()
        conn.execute("INSERT INTO upgrade_requests (team_id,username,current_plan,requested_plan,notes) VALUES (?,?,?,?,?)",
                     (my_team, me.get("username",""), org_plan, requested_plan, notes.strip()))
        conn.commit(); conn.close()
        log_audit(my_team, me.get("username",""), my_role, "upgrade.request", "org", my_team, f"{org_plan}->{requested_plan}")
        st.toast("Upgrade request submitted ✅", icon="✅")

    st.markdown("---")
    st.subheader("Org Workspace")
    render_team_intel(prefix="orgadmin")

def render_root_admin():
    st.header("🛡 Root Admin")
    st.caption("SaaS owner backend: orgs, users, credits, upgrades, health, logs.")

    tabs = st.tabs(["🏢 Orgs", "👥 Users", "💳 Credits", "⬆ Upgrades", "🩺 SaaS Health", "📜 Global Logs"])

    with tabs[0]:
        conn = db_conn()
        odf = pd.read_sql_query("SELECT team_id,org_name,plan,seats_allowed,status,allowed_agents_json,created_at FROM orgs ORDER BY created_at DESC", conn)
        conn.close()
        st.dataframe(odf, use_container_width=True, hide_index=True)

    with tabs[1]:
        conn = db_conn()
        udf = pd.read_sql_query("SELECT username,name,email,role,credits,active,team_id,created_at,last_login_at FROM users ORDER BY created_at DESC", conn)
        conn.close()
        st.dataframe(udf, use_container_width=True, hide_index=True)

        with st.form("root_add_user"):
            st.subheader("Add User (manual)")
            team_id = st.text_input("Team ID")
            username = st.text_input("Username")
            name = st.text_input("Name")
            email = st.text_input("Email")
            role = st.selectbox("Role", ["viewer","editor","admin"], index=0)
            pw = st.text_input("Password", type="password")
            submit = st.form_submit_button("Create User", use_container_width=True)
        if submit and team_id.strip() and username.strip() and pw.strip():
            conn = db_conn()
            conn.execute("INSERT INTO users (username,email,name,password,role,active,plan,credits,verified,team_id) VALUES (?,?,?,?,?,1,?,?,1,?)",
                         (username.strip(),email.strip(),name.strip(),_hash_password(pw.strip()),role,get_org(team_id.strip()).get("plan","Lite"),10,team_id.strip()))
            conn.commit(); conn.close()
            st.toast("User created ✅", icon="✅")
            st.rerun()

        with st.form("root_remove_user"):
            st.subheader("Remove User")
            del_user = st.text_input("Username to delete")
            submit = st.form_submit_button("Delete", use_container_width=True)
        if submit and del_user.strip():
            conn = db_conn()
            conn.execute("DELETE FROM users WHERE username=? AND role!='root'", (del_user.strip(),))
            conn.commit(); conn.close()
            st.toast("User deleted ✅", icon="✅")
            st.rerun()

    with tabs[2]:
        st.subheader("Credits")
        with st.form("root_credits"):
            username = st.text_input("Username")
            delta = st.number_input("Add credits (+/-)", value=10, step=1)
            submit = st.form_submit_button("Apply", use_container_width=True)
        if submit and username.strip():
            conn = db_conn()
            conn.execute("UPDATE users SET credits=COALESCE(credits,0)+? WHERE username=?", (int(delta), username.strip()))
            conn.commit(); conn.close()
            st.toast("Credits updated ✅", icon="✅")

    with tabs[3]:
        st.subheader("Upgrade Requests")
        conn = db_conn()
        rdf = pd.read_sql_query("SELECT id,team_id,username,current_plan,requested_plan,notes,status,created_at FROM upgrade_requests ORDER BY id DESC", conn)
        conn.close()
        st.dataframe(rdf, use_container_width=True, hide_index=True)

        with st.form("approve_upgrade"):
            req_id = st.number_input("Request ID", min_value=1, step=1)
            approve = st.selectbox("Action", ["approve", "reject"], index=0)
            submit = st.form_submit_button("Apply", use_container_width=True)
        if submit:
            conn = db_conn()
            df = pd.read_sql_query("SELECT * FROM upgrade_requests WHERE id=?", conn, params=(int(req_id),))
            if df.empty:
                conn.close()
                st.error("Request not found.")
            else:
                row = df.iloc[0].to_dict()
                if approve == "approve":
                    # apply plan + auto agents
                    plan = row["requested_plan"]
                    seats = PLAN_SEATS.get(plan, 1)
                    agents = default_allowed_agents_for_plan(plan)
                    conn.execute("UPDATE orgs SET plan=?, seats_allowed=?, allowed_agents_json=? WHERE team_id=?",
                                 (plan, int(seats), json.dumps(agents), row["team_id"]))
                    conn.execute("UPDATE upgrade_requests SET status='approved' WHERE id=?", (int(req_id),))
                else:
                    conn.execute("UPDATE upgrade_requests SET status='rejected' WHERE id=?", (int(req_id),))
                conn.commit(); conn.close()
                st.toast("Upgrade processed ✅", icon="✅")
                st.rerun()

    with tabs[4]:
        st.subheader("SaaS Health")
        st.markdown("<div class='ms-card'><b>DB:</b> breatheeasy.db • <b>UTC:</b> "
                    f"{datetime.utcnow().isoformat()} • <b>Python:</b> {os.sys.version.split()[0]}</div>",
                    unsafe_allow_html=True)
        conn = db_conn()
        tables = pd.read_sql_query("SELECT name FROM sqlite_master WHERE type='table' ORDER BY name", conn)
        conn.close()
        st.dataframe(tables, use_container_width=True, hide_index=True)

    with tabs[5]:
        conn = db_conn()
        gdf = pd.read_sql_query("SELECT timestamp,team_id,actor,actor_role,action_type,object_type,object_id,details FROM audit_logs ORDER BY id DESC LIMIT 500", conn)
        conn.close()
        st.dataframe(gdf, use_container_width=True, hide_index=True)

# ============================================================
# PAGES
# ============================================================
def render_dashboard():
    st.markdown(f"<div class='ms-card'><h2 style='margin:0;'>🏠 Dashboard</h2>"
                f"<div class='ms-muted'>Team `{my_team}` • Plan `{org_plan}`</div></div>",
                unsafe_allow_html=True)

    if st.session_state["swarm_running"]:
        q = st.session_state["swarm_queue"]
        idx = int(st.session_state["swarm_idx"])
        nxt = q[idx] if idx < len(q) else "—"
        st.markdown(f"<div class='ms-card'><b>🚀 Swarm Running</b> <span class='ms-chip'>Next: {nxt}</span>"
                    f"<div class='ms-muted'>Auto-run delay: {st.session_state['swarm_autodelay']}s • Auto-run: {st.session_state['swarm_autorun']}</div></div>",
                    unsafe_allow_html=True)
    else:
        st.markdown("<div class='ms-card'><b>Swarm Status:</b> Idle</div>", unsafe_allow_html=True)

    selected = st.session_state.get("last_active_swarm", []) or []
    rep = st.session_state.get("report", {}) or {}
    if selected:
        st.subheader("Report Integrity Check")
        df = report_integrity_df(rep, selected)
        st.dataframe(df, use_container_width=True, hide_index=True)
        empty = [r["agent"] for r in df.to_dict("records") if r["status"] == "EMPTY"]
        if empty:
            st.warning(f"Empty outputs: {empty}")
            cols = st.columns(4)
            for i, a in enumerate(empty):
                with cols[i % 4]:
                    st.button(f"Retry {a}", key=f"dash_retry_{a}", on_click=retry_agent, args=(a,), use_container_width=True)
    else:
        st.info("Run a swarm to see integrity results.")

def render_guide():
    st.header(f"📖 {APP_NAME} Guide")
    st.markdown("### Protocol")
    for line in DEPLOY_PROTOCOL:
        st.markdown(f"- {line}")

def seat_how_to_use(agent_key: str) -> str:
    guides = {
        "analyst": "Use to set pricing & offers based on competitor gaps.",
        "marketing_adviser": "Use to choose channel priorities + messaging pillars + KPIs.",
        "market_researcher": "Use to refine ICP/segments + demand themes.",
        "ecommerce_marketer": "Use to implement offer ladder + flows + remarketing.",
        "ads": "Paste into Google/Meta. Replace claims with proof points.",
        "creative": "Use as a design brief + prompt pack.",
        "guest_posting": "Use outreach templates; track in Projects.",
        "strategist": "Convert roadmap into tasks; review weekly.",
        "social": "Schedule the calendar; recycle winning hooks.",
        "geo": "Apply GBP checklist + citations + review system.",
        "gbp_growth": "Post weekly on GBP, reply to reviews, triage ranking drops.",
        "audit": "Fix top friction; rerun after changes.",
        "seo": "Publish pillar + build cluster pages.",
    }
    return guides.get(agent_key, "Apply into execution.")

# ============================================================
# ROUTER (Quick Nav)
# ============================================================
def resolve_page():
    choice = st.session_state.get("nav_choice", "🏠 Dashboard")
    if choice == "🏠 Dashboard":
        render_dashboard()
        return
    if choice == "📖 Guide":
        render_guide()
        return
    if choice == "🤝 Team Intel":
        render_team_intel(prefix="teamintel")
        return
    if choice == "⚙ Org Admin":
        render_org_admin()
        return
    if choice == "🛡 Root Admin":
        render_root_admin()
        return

    # Agent seats
    for lbl, key in AGENT_UI:
        if choice == lbl:
            render_seat(lbl, key)
            return

# ============================================================
# MAIN NAV (tabs for visual navigation + quick nav actually renders)
# ============================================================
# Keep tabs for UI, but we render content via Quick Nav to avoid duplicate keys/forms.
tab_labels = ["🏠 Dashboard", "📖 Guide"] + [lbl for lbl, _ in AGENT_UI] + ["🤝 Team Intel"]
if my_role in {"admin", "root"}:
    tab_labels.append("⚙ Org Admin")
if is_root:
    tab_labels.append("🛡 Root Admin")

# Visual tabs (do not duplicate forms; just set nav_choice)
tabs = st.tabs(tab_labels)
for i, name in enumerate(tab_labels):
    with tabs[i]:
        if st.button(f"Open {name}", key=f"open_{i}_{name}", use_container_width=True):
            st.session_state["nav_choice"] = name
            st.rerun()

# Render chosen page once
resolve_page()
